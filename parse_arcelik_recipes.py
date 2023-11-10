import re
import os
import json
import regex

import docx


from calc_recipe_ingredient_info_distances import get_graph_nodes

def get_text(file_name):
    document = docx.Document(file_name)
    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)

    return "\n".join(full_text)

def parse_doc_to_recipe(file_name):
    recipe = {"instructions": [], "original_ingredients": []}
    unused_lines = []

    is_iterating_through_ingredients_block = False
    is_iterating_through_instructions_block = False

    document = docx.Document(file_name)
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if not paragraph_text or "~~" in paragraph_text:
            continue
        if "title" not in recipe:
            recipe["title"] = paragraph_text
            continue
        if re.match(r"ingredients", paragraph_text.lower()):
            is_iterating_through_ingredients_block = True
            is_iterating_through_instructions_block = False
            continue
        if re.match(r"instructions", paragraph_text.lower()):
            is_iterating_through_instructions_block = True
            is_iterating_through_ingredients_block = False
            continue
        if is_iterating_through_ingredients_block:
            recipe["original_ingredients"].append(paragraph_text)
        if is_iterating_through_instructions_block:
            recipe["instructions"].append(paragraph_text)
        else:
            unused_lines.append(paragraph_text)

    return recipe

def parse_ingredients(recipe: list, ingr_labels_raw_to_r1m: dict, ingr_labels_r1m_to_raw: dict) -> tuple:
    regex_measurements_without_fractions = r"\d+(\.\d+)?\s*(?:fluid\s*|fl\.?\s*|dry\s*)?(?:(?:teaspoons?|tsps?)|t\.?\s*sp\.?\s*|tablespoons?|tbsps?|tbsp\.?|cups?|c\.?\s*|ounces?|oz\.?|pounds?|lbs?|grams?|g\.?|kilograms?|kg\.?|milliliters?|ml\.?|liters?|l\.?|quarts?|qts?|gallons?|gals?|pinch(?:es)?|dashes?)\s*(?:(?:of|in|for)\s*)?"
    regex_measurements_with_fractions = r"\d+(\s*\d+/\d+)?\s*(?:fluid\s*|fl\.?\s*|dry\s*)?(?:(?:teaspoons?|tsps?)|t\.?\s*sp\.?\s*|tablespoons?|tbsps?|tbsp\.?|cups?|c\.?\s*|ounces?|oz\.?|pounds?|lbs?|grams?|g\.?|kilograms?|kg\.?|milliliters?|ml\.?|liters?|l\.?|quarts?|qts?|gallons?|gals?|pinch(?:es)?|dashes?)\s*(?:(?:of|in|for)\s*)?"


    ingredient_label_agg_map = {
        "egg yolks-120 g yumurta sar\u0131s\u0131": "egg",
        "egg yolks-": "egg",
        "egg": "egg",
        "eggs-4 yumurta beyaz\u0131": "egg",
        "medium eggs, beaten": "egg",
        "egg-1 yumurta": "egg",
        "egg+2 tbsp milk": "egg",
        "egg white": "egg",
        "egg whites": "egg",
        "egg yolk": "egg",
        "egg yolks": "egg",
        "medium eggs-": "egg",
        "egg white, lightly beaten": "egg",
        "large eggs, beaten": "egg",
        "large eggs, separated": "egg",
        "egg, beaten": "egg",
        "medium egg": "egg",
        "eggs": "egg",
        "milk-": "milk",
        "milk": "milk",
        "milk 125 g": "milk",
        "milk-250 ml s\u00fct": "milk",
        "milk-30 g s\u00fct": "milk",
        "teaspoon salt": "salt",
        "of salt-bir tutam tuz": "salt",
        "salt, pepper - tuz biber": ["salt", "pepper"],
        "a pinch of salt": "salt",
        "salt, crushed black pepper flakes- tuz, ezilmi\u015f karabiber": ["salt", "pepper"],
        "salt, pepper and nutmeg": ["salt", "pepper", "nutmeg"],
        "#salt and freshly ground black pepper": ["salt", "pepper"],
        "salt and ground pepper": ["salt", "pepper"],
        "salt+pepper": ["salt", "pepper"],
        "of pink salt": "salt",
        "salt, black pepper": ["salt", "pepper"],
        "salt pepper": ["salt", "pepper"],
        "salt and pepper to taste": ["salt", "pepper"],
        "salt and freshly ground black pepper": ["salt", "pepper"],
        "of salt": "salt",
        "salt and freshly ground black pepper-": ["salt", "pepper"],
        "salt and freshly ground pepper": ["salt", "pepper"],
        "salt and pepper": ["salt", "pepper"],
        "salt-1 \u00e7ay ka\u015f\u0131\u011f\u0131 tuz": "salt",
        "salt and freshly ground pepper to season": ["salt", "pepper"],
        "pepper and salt, optional": ["salt", "pepper"],
        "salt, pepper": ["salt", "pepper"],
        "flaked sea salt": "salt",
        "salt-\u00bd g tuz": "salt",
        "of garlic": "garlic",
        "of garlic -chopped": "garlic",
        "garlic cloves, peeled and crushed": "garlic",
        "garlic cloves, crushed": "garlic",
        "of garlic-chopped": "garlic",
        "garlic cloves, minced": "garlic",
        "garlic minced": "garlic",
        "garlic cloves, cut into slivers": "garlic",
        "of garlic- chopped": "garlic",
        "garlic cloves": "garlic",
        "garlic clove, crushed": "garlic",
        "garlic": "garlic",
        "garlic clove,  finely chopped": "garlic",
        "whole garlic-halved": "garlic",
        "coconut sugar": "sugar",
        "demerara sugar": "sugar",
        "sugar": "sugar",
        "caster sugar": "sugar",
        "sugar -100 g \u015feker": "sugar",
        "granulated white sugar": "sugar",
        "coconut sugar-175 g esmer \u015feker": "sugar",
        "coconut sugar-75 g  esmer \u015feker": "sugar",
        "granulated sugar": "sugar",
        "brown sugar": "sugar",
        "sugar-150 g (\u00be su barda\u011f\u0131 )\u015feker": "sugar",
        "icing sugar": "sugar",
        "dememera sugar": "sugar",
        "powdered sugar 50 g": "sugar",
        "juice of a lemon": "lemon",
        "lemon zest": "lemon",
        "fresh lemon\u00a0juice": "lemon",
        "lemon,": "lemon",
        "juice of lemon": "lemon",
        "lemon juice": "lemon",
        "juice of  of a lemon": "lemon",
        "finely grated zest of 1 lemon": "lemon",
        "zest of a lemon": "lemon",
        "fresh lemon juice": "lemon",
        "lemon": "lemon",
        "juice of a fresh lemon + 1 sliced lemon": "lemon",
        "lemon, zest and juice": "lemon",
        "lemon cut into quarter": "lemon",
        "juice of half a lemon": "lemon",
        "juice of 2 lemons + 1 lemon, cut into 4": "lemon",
        "flour to roll": "flour",
        "cornflour": "flour",
        "oat flour": "flour",
        "whole wheat flour with pinch of salt": "flour",
        "whole wheat flour": "flour",
        "strong whole wheat flour-500 gr g\u00fc\u00e7l\u00fc tam bu\u011fday unu": "flour",
        "strong whole wheat bread flour": "flour",
        "all-purpose flour": "flour",
        "whole wheat flour-": "flour",
        "flour": "flour",
        "corn flour": "flour",
        "red onions-cut into quarter": "onion",
        "onion,": "onion",
        "large onions, peeled and diced": "onion",
        "pureed onion": "onion",
        "onion": "onion",
        "onions": "onion",
        "green onions": "onion",
        "small onion, peeled and diced": "onion",
        "red onions cut into quarters": "onion",
        "onion-finely chopped": "onion",
        "onion, diced": "onion",
        "parsley, chopped": "parsley",
        "fresh parsley, finely chopped": "parsley",
        "finely chopped parsley": "parsley",
        "fresh parsley leaves": "parsley",
        "fresh parsley": "parsley",
        "fresh flat leaf parsley leaves": "parsley",
        "fresh parsley, chopped": "parsley",
        "parsley,": "parsley",
        "leaves of fresh basil": "basil",
        "fresh basil": "basil",
        "leaves fresh basil": "basil",
        "basil pesto-1/3 fincan fesle\u011fen pesto": "basil",
        "sun dried tomatoes, chopped": "tomato",
        "tomatoes-grated": "tomato",
        "tomatoes": "tomato",
        "tomato puree": "tomato",
        "tomatoes, chopped": "tomato",
        "can crushed tomatoes": "tomato",
        "tomatoes, crushed": "tomato",
        "tomatoes,": "tomato",
        "tomatoes- chopped": "tomato",
        "vanilla extract\n-50 g": "vanilla",
        "vanilla": "vanilla",
        "vanilla bean-1 vanilya fasulyesi": "vanilla",
        "vanilla extract-5 ml vanilya \u00f6z\u00fc": "vanilla",
        "vanilla pod": "vanilla",
        "vanilla essence": "vanilla",
        "vanilla extract": "vanilla",
        "vanilla bean (or 1 tsp vanilla extract)": "vanilla",
        "heavy cream-250g a\u011f\u0131r krema": "cream",
        "single cream": "cream",
        "heavy cream 50 g": "cream",
        "cream, room temperature": "cream",
        "cream cheese 100 g": "cream",
        "cream cheese, room temperature": "cream",
        "coriander leaves": "coriander",
        "fresh coriander, chopped": "coriander",
        "fresh coriander leaves": "coriander",
        "chilli pepper": "chilli",
        "small chilli,de-seeded & finely chopped": "chilli",
        "chilli powder": "chilli",
        "chilli flakes": "chilli",
        "fresh chopped thyme": "thyme",
        "spring of thyme": "thyme",
        "of fresh thyme": "thyme",
        "fresh thyme and fresh tarragon": "thyme",
        "chopped fresh herbs (thyme, marjoram, sage)": "thyme",
        "thyme(optional)": "thyme",
        "dried thyme": "thyme",
        "fresh thyme and bay leaves": "thyme",
        "fresh thyme": "thyme",
        "of toasted almonds": "almond",
        "almonds,": "almond",
        "almonds-chopped-100 gr badem do\u011franm\u0131\u015f": "almond",
        "ground almonds": "almond",
        "oatmeal-": "oatmeal",
        "rolled oatmeal": "oatmeal",
        "oatmeal": "oatmeal",
        "active dried yeast  or 15g/1/2 oz fresh yeast": "yeast",
        "sachet dried yeast or 15g fresh yeast-7 gr po\u015fet kuru maya veya 15 gr taze maya": "yeast",
        "active dry yeast": "yeast",
        "active dried yeast  or 15g  fresh yeast": "yeast",
        "dried yeast or 15g fresh yeast": "yeast",
        "fresh yeast": "yeast",
        "balsamic vinegar": "vinegar",
        "vinegar": "vinegar",
    }

    recipe["quantities"] = []
    recipe["ingredients"] = []
    unmatched_ingredient_labels = []

    for ingredient in recipe["original_ingredients"]:
        ingredient_parts_raw = ingredient.split(";")
        measurement_raw = ingredient_parts_raw[0].lower().strip()
        ingredient_label_raw = ingredient_parts_raw[1].lower().strip()

        if ingredient_label_raw in ingredient_label_agg_map:
            mapped_items = ingredient_label_agg_map[ingredient_label_raw]
            if isinstance(mapped_items, list):
                for mapped_item in mapped_items:
                    unmatched_ingredient_labels.append(mapped_item)
                    ingredient_label_raw = mapped_item
            else:
                ingredient_label_raw = mapped_items

        # handle ingredient labels
        if ingredient_label_raw not in ingr_labels_raw_to_r1m:
            unmatched_ingredient_labels.append(ingredient_label_raw)
            continue

        # TODO append the parsed ingredient list to the reicpe object


        # handle quantities
        measurement = regex.match(regex_measurements_with_fractions, ingredient)
        test1 = measurement.groups() if measurement else None
        # test = measurement[0]
        measurement2 = regex.match(regex_measurements_without_fractions, ingredient)
        # test2 = measurement[0]
        test2 = measurement2.groups() if measurement2 else None
        m1 = regex.search(regex_measurements_with_fractions, ingredient)
        test3 = m1.groups() if m1 else None
        m2 = regex.search(regex_measurements_without_fractions, ingredient)
        test4 = m2.groups() if m2 else None
        # TODO append the quantity to the recipe object

    return recipe, unmatched_ingredient_labels

def main():
    ARCELIC_RECIPES_DIR_PATH = os.path.abspath("./inputs/arcelik_recipe/Revize Edilen Tarifler/")
    GRAPH_NODES_PATH = os.path.abspath("./inputs/graph/nodes_191120.csv")

    unmatched_ingredient_labels = []

    fg_ingredients = get_graph_nodes(GRAPH_NODES_PATH)

    # TODO load from json
    ingr_labels_raw_to_r1m = {"salt": ["salt", "rock_salt"]}
    ingr_labels_r1m_to_raw = {"salt": ["salt"],
                              "rock_salt": ["salt"]}

    for file_name in os.listdir(ARCELIC_RECIPES_DIR_PATH):
        file_path = os.path.join(ARCELIC_RECIPES_DIR_PATH, file_name)
        recipe = parse_doc_to_recipe(file_path)
        recipe, _unmatched_ingredient_labels = parse_ingredients(recipe, ingr_labels_raw_to_r1m, ingr_labels_r1m_to_raw)
        unmatched_ingredient_labels += _unmatched_ingredient_labels

    print(len(unmatched_ingredient_labels))
    unmatched_ingredient_labels = list(set(unmatched_ingredient_labels))
    print(len(unmatched_ingredient_labels))


    sugar = []
    lemon = []
    flour = []
    onion = []
    parsley = []
    basil = []
    tomato = []
    vanilla = []
    chilli = []
    coriander = []
    oatmeal = []
    cream = []
    thyme = []
    almond = []
    yeast = []
    vinegar = []

    for unmatched_ingredient in unmatched_ingredient_labels:
        if "sugar" in unmatched_ingredient:
            sugar.append(unmatched_ingredient)
        if "lemon" in unmatched_ingredient:
            lemon.append(unmatched_ingredient)
        if "flour" in unmatched_ingredient:
            flour.append(unmatched_ingredient)
        if "onion" in unmatched_ingredient:
            onion.append(unmatched_ingredient)
        if "parsley" in unmatched_ingredient:
            parsley.append(unmatched_ingredient)
        if "basil" in unmatched_ingredient:
            basil.append(unmatched_ingredient)
        if "tomato" in unmatched_ingredient:
            tomato.append(unmatched_ingredient)
        if "vanilla" in unmatched_ingredient:
            vanilla.append(unmatched_ingredient)
        if "cream" in unmatched_ingredient:
            cream.append(unmatched_ingredient)
        if "coriander" in unmatched_ingredient:
            coriander.append(unmatched_ingredient)
        if "chilli" in unmatched_ingredient:
            chilli.append(unmatched_ingredient)
        if "thyme" in unmatched_ingredient:
            thyme.append(unmatched_ingredient)
        if "almond" in unmatched_ingredient:
            almond.append(unmatched_ingredient)
        if "oatmeal" in unmatched_ingredient:
            oatmeal.append(unmatched_ingredient)
        if "yeast" in unmatched_ingredient:
            yeast.append(unmatched_ingredient)
        if "vinegar" in unmatched_ingredient:
            vinegar.append(unmatched_ingredient)

    unmatched_ingredient_labels.append(sugar)
    unmatched_ingredient_labels.append(lemon)
    unmatched_ingredient_labels.append(flour)
    unmatched_ingredient_labels.append(onion)
    unmatched_ingredient_labels.append(parsley)
    unmatched_ingredient_labels.append(basil)
    unmatched_ingredient_labels.append(tomato)
    unmatched_ingredient_labels.append(vanilla)
    unmatched_ingredient_labels.append(cream)
    unmatched_ingredient_labels.append(coriander)
    unmatched_ingredient_labels.append(chilli)
    unmatched_ingredient_labels.append(thyme)
    unmatched_ingredient_labels.append(almond)
    unmatched_ingredient_labels.append(oatmeal)
    unmatched_ingredient_labels.append(yeast)
    unmatched_ingredient_labels.append(vinegar)

    with open(os.path.abspath("./outputs/arcelic_raw_ingredient_labels.json"), "w") as json_file:
        json.dump(unmatched_ingredient_labels, json_file, indent=2)

if __name__ == "__main__":
    main()