import re
import os
import json
import regex

import spacy
import docx

from nltk.stem import PorterStemmer
from calc_recipe_ingredient_info_distances import get_graph_nodes

# initialize the nlp stuff
nlp = spacy.load("en_core_web_sm")
stemmer = PorterStemmer()

def get_text(file_name):
    document = docx.Document(file_name)
    full_text = []
    for paragraph in document.paragraphs:
        full_text.append(paragraph.text)

    return "\n".join(full_text)

def parse_doc_to_recipe(file_name):
    recipe = {"instructions": [], "original_ingredients": []}
    unused_lines = []

    is_iterating_through_ingredients_block = False
    is_iterating_through_instructions_block = False

    document = docx.Document(file_name)
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if not paragraph_text or "~~" in paragraph_text:
            continue
        if "title" not in recipe:
            recipe["title"] = paragraph_text
            continue
        if re.match(r"ingredients", paragraph_text.lower()):
            is_iterating_through_ingredients_block = True
            is_iterating_through_instructions_block = False
            continue
        if re.match(r"instructions", paragraph_text.lower()):
            is_iterating_through_instructions_block = True
            is_iterating_through_ingredients_block = False
            continue
        if is_iterating_through_ingredients_block:
            recipe["original_ingredients"].append(paragraph_text)
        if is_iterating_through_instructions_block:
            recipe["instructions"].append(paragraph_text)
        else:
            unused_lines.append(paragraph_text)

    return recipe

# Function to extract the normalized main keyword in singular form (as suggested by chatgippity)
def extract_main_keyword_singular(ingredient):
    # Process the ingredient text using spaCy
    doc = nlp(ingredient)

    # Extract the main noun or entity
    main_keywords = [token.text.lower() for token in doc if token.pos_ == "NOUN"]

    # If no noun is found, consider the whole ingredient text
    if not main_keywords:
        main_keywords = [token.text.lower() for token in doc]

    # Perform stemming to get the singular form
    singular_keywords = [stemmer.stem(word) for word in main_keywords]

    # Join the words to get the normalized main keyword in singular form
    normalized_main_keyword_singular = " ".join(singular_keywords)

    return normalized_main_keyword_singular

def executeHardcodedIngredientMappings(unmatched_ingredient_labels):
    sugar = []
    lemon = []
    flour = []
    onion = []
    parsley = []
    basil = []
    tomato = []
    vanilla = []
    chilli = []
    coriander = []
    oatmeal = []
    cream = []
    thyme = []
    almond = []
    yeast = []
    vinegar = []
    beef = []
    pork = []
    baking = []
    turkey = []
    orange = []
    potato = []
    lime = []
    chicken = []
    mint = []
    rosemary = []
    spinach = []
    nutmeg = []

    for unmatched_ingredient in unmatched_ingredient_labels:
        if "sugar" in unmatched_ingredient:
            sugar.append(unmatched_ingredient)
        if "lemon" in unmatched_ingredient:
            lemon.append(unmatched_ingredient)
        if "flour" in unmatched_ingredient:
            flour.append(unmatched_ingredient)
        if "onion" in unmatched_ingredient:
            onion.append(unmatched_ingredient)
        if "parsley" in unmatched_ingredient:
            parsley.append(unmatched_ingredient)
        if "basil" in unmatched_ingredient:
            basil.append(unmatched_ingredient)
        if "tomato" in unmatched_ingredient:
            tomato.append(unmatched_ingredient)
        if "vanilla" in unmatched_ingredient:
            vanilla.append(unmatched_ingredient)
        if "cream" in unmatched_ingredient:
            cream.append(unmatched_ingredient)
        if "coriander" in unmatched_ingredient:
            coriander.append(unmatched_ingredient)
        if "chilli" in unmatched_ingredient:
            chilli.append(unmatched_ingredient)
        if "thyme" in unmatched_ingredient:
            thyme.append(unmatched_ingredient)
        if "almond" in unmatched_ingredient:
            almond.append(unmatched_ingredient)
        if "oatmeal" in unmatched_ingredient:
            oatmeal.append(unmatched_ingredient)
        if "yeast" in unmatched_ingredient:
            yeast.append(unmatched_ingredient)
        if "vinegar" in unmatched_ingredient:
            vinegar.append(unmatched_ingredient)

        if "beef" in unmatched_ingredient:
            beef.append(unmatched_ingredient)
        if "pork" in unmatched_ingredient:
            pork.append(unmatched_ingredient)
        if "baking" in unmatched_ingredient:
            baking.append(unmatched_ingredient)
        if "turkey" in unmatched_ingredient:
            turkey.append(unmatched_ingredient)
        if "orange" in unmatched_ingredient:
            orange.append(unmatched_ingredient)
        if "potato" in unmatched_ingredient:
            potato.append(unmatched_ingredient)
        if "lime" in unmatched_ingredient:
            lime.append(unmatched_ingredient)
        if "chicken" in unmatched_ingredient:
            chicken.append(unmatched_ingredient)
        if "mint" in unmatched_ingredient:
            mint.append(unmatched_ingredient)
        if "rosemary" in unmatched_ingredient:
            rosemary.append(unmatched_ingredient)
        if "spinach" in unmatched_ingredient:
            spinach.append(unmatched_ingredient)
        if "nutmeg" in unmatched_ingredient:
            nutmeg.append(unmatched_ingredient)

    # unmatched_ingredient_labels.append(sugar)
    # unmatched_ingredient_labels.append(lemon)
    # unmatched_ingredient_labels.append(flour)
    # unmatched_ingredient_labels.append(onion)
    # unmatched_ingredient_labels.append(parsley)
    # unmatched_ingredient_labels.append(basil)
    # unmatched_ingredient_labels.append(tomato)
    # unmatched_ingredient_labels.append(vanilla)
    # unmatched_ingredient_labels.append(cream)
    # unmatched_ingredient_labels.append(coriander)
    # unmatched_ingredient_labels.append(chilli)
    # unmatched_ingredient_labels.append(thyme)
    # unmatched_ingredient_labels.append(almond)
    # unmatched_ingredient_labels.append(oatmeal)
    # unmatched_ingredient_labels.append(yeast)
    # unmatched_ingredient_labels.append(vinegar)

    # unmatched_ingredient_labels.append(beef)
    # unmatched_ingredient_labels.append(pork)
    # unmatched_ingredient_labels.append(baking)
    # unmatched_ingredient_labels.append(turkey)
    # unmatched_ingredient_labels.append(orange)
    # unmatched_ingredient_labels.append(potato)
    # unmatched_ingredient_labels.append(lime)
    # unmatched_ingredient_labels.append(chicken)
    # unmatched_ingredient_labels.append(mint)
    # unmatched_ingredient_labels.append(rosemary)
    # unmatched_ingredient_labels.append(spinach)
    # unmatched_ingredient_labels.append(nutmeg)

    return unmatched_ingredient_labels

def parse_ingredients(recipe: list, ingr_labels_raw_to_r1m: dict, ingr_labels_r1m_to_raw: dict) -> tuple:
    regex_measurements_without_fractions = r"\d+(\.\d+)?\s*(?:fluid\s*|fl\.?\s*|dry\s*)?(?:(?:teaspoons?|tsps?)|t\.?\s*sp\.?\s*|tablespoons?|tbsps?|tbsp\.?|cups?|c\.?\s*|ounces?|oz\.?|pounds?|lbs?|grams?|g\.?|kilograms?|kg\.?|milliliters?|ml\.?|liters?|l\.?|quarts?|qts?|gallons?|gals?|pinch(?:es)?|dashes?)\s*(?:(?:of|in|for)\s*)?"
    regex_measurements_with_fractions = r"\d+(\s*\d+/\d+)?\s*(?:fluid\s*|fl\.?\s*|dry\s*)?(?:(?:teaspoons?|tsps?)|t\.?\s*sp\.?\s*|tablespoons?|tbsps?|tbsp\.?|cups?|c\.?\s*|ounces?|oz\.?|pounds?|lbs?|grams?|g\.?|kilograms?|kg\.?|milliliters?|ml\.?|liters?|l\.?|quarts?|qts?|gallons?|gals?|pinch(?:es)?|dashes?)\s*(?:(?:of|in|for)\s*)?"


    ingredient_label_agg_map = {
        "egg yolks-120 g yumurta sar\u0131s\u0131": "egg",
        "egg yolks-": "egg",
        "egg": "egg",
        "eggs-4 yumurta beyaz\u0131": "egg",
        "medium eggs, beaten": "egg",
        "egg-1 yumurta": "egg",
        "egg+2 tbsp milk": "egg",
        "egg white": "egg",
        "egg whites": "egg",
        "egg yolk": "egg",
        "egg yolks": "egg",
        "medium eggs-": "egg",
        "egg white, lightly beaten": "egg",
        "eggs\u00a0(at room temperature)": "egg",
        "large eggs, beaten": "egg",
        "large eggs, separated": "egg",
        "egg, beaten": "egg",
        "medium egg": "egg",
        "eggs": "egg",
        "milk-": "milk",
        "milk": "milk",
        "milk 125 g": "milk",
        "milk-250 ml s\u00fct": "milk",
        "milk-30 g s\u00fct": "milk",
        "teaspoon salt": "salt",
        "of salt-bir tutam tuz": "salt",
        "salt, pepper - tuz biber": ["salt", "pepper"],
        "a pinch of salt": "salt",
        "salt, crushed black pepper flakes- tuz, ezilmi\u015f karabiber": ["salt", "pepper"],
        "salt, pepper and nutmeg": ["salt", "pepper", "nutmeg"],
        "#salt and freshly ground black pepper": ["salt", "pepper"],
        "salt and ground pepper": ["salt", "pepper"],
        "salt+pepper": ["salt", "pepper"],
        "of pink salt": "salt",
        "salt, black pepper": ["salt", "pepper"],
        "salt pepper": ["salt", "pepper"],
        "salt and pepper to taste": ["salt", "pepper"],
        "salt and freshly ground black pepper": ["salt", "pepper"],
        "of salt": "salt",
        "salt and freshly ground black pepper-": ["salt", "pepper"],
        "salt and freshly ground pepper": ["salt", "pepper"],
        "salt and pepper": ["salt", "pepper"],
        "salt-1 \u00e7ay ka\u015f\u0131\u011f\u0131 tuz": "salt",
        "salt and freshly ground pepper to season": ["salt", "pepper"],
        "pepper and salt, optional": ["salt", "pepper"],
        "salt, pepper": ["salt", "pepper"],
        "flaked sea salt": "salt",
        "salt-\u00bd g tuz": "salt",
        "of garlic": "garlic",
        "of garlic -chopped": "garlic",
        "garlic cloves, peeled and crushed": "garlic",
        "garlic cloves, crushed": "garlic",
        "of garlic-chopped": "garlic",
        "garlic cloves, minced": "garlic",
        "garlic minced": "garlic",
        "garlic cloves, cut into slivers": "garlic",
        "of garlic- chopped": "garlic",
        "garlic cloves": "garlic",
        "garlic clove, crushed": "garlic",
        "garlic": "garlic",
        "garlic clove,  finely chopped": "garlic",
        "whole garlic-halved": "garlic",
        "coconut sugar": "sugar",
        "demerara sugar": "sugar",
        "sugar": "sugar",
        "caster sugar": "sugar",
        "sugar -100 g \u015feker": "sugar",
        "granulated white sugar": "sugar",
        "coconut sugar-175 g esmer \u015feker": "sugar",
        "coconut sugar-75 g  esmer \u015feker": "sugar",
        "granulated sugar": "sugar",
        "brown sugar": "sugar",
        "sugar-150 g (\u00be su barda\u011f\u0131 )\u015feker": "sugar",
        "icing sugar": "sugar",
        "dememera sugar": "sugar",
        "powdered sugar 50 g": "sugar",
        "juice of a lemon": "lemon",
        "lemon zest": "lemon",
        "fresh lemon\u00a0juice": "lemon",
        "lemon,": "lemon",
        "juice of lemon": "lemon",
        "lemon juice": "lemon",
        "juice of  of a lemon": "lemon",
        "finely grated zest of 1 lemon": "lemon",
        "zest of a lemon": "lemon",
        "fresh lemon juice": "lemon",
        "lemon": "lemon",
        "juice of a fresh lemon + 1 sliced lemon": "lemon",
        "lemon, zest and juice": "lemon",
        "lemon cut into quarter": "lemon",
        "juice of half a lemon": "lemon",
        "juice of 2 lemons + 1 lemon, cut into 4": "lemon",
        "flour to roll": "flour",
        "cornflour": "flour",
        "oat flour": "flour",
        "whole wheat flour with pinch of salt": "flour",
        "whole wheat flour": "flour",
        "strong whole wheat flour-500 gr g\u00fc\u00e7l\u00fc tam bu\u011fday unu": "flour",
        "strong whole wheat bread flour": "flour",
        "all-purpose flour": "flour",
        "whole wheat flour-": "flour",
        "flour": "flour",
        "corn flour": "flour",
        "red onions-cut into quarter": "onion",
        "onion,": "onion",
        "large onions, peeled and diced": "onion",
        "pureed onion": "onion",
        "onion": "onion",
        "onions": "onion",
        "green onions": "onion",
        "small onion, peeled and diced": "onion",
        "red onions cut into quarters": "onion",
        "onion-finely chopped": "onion",
        "onion, diced": "onion",
        "parsley, chopped": "parsley",
        "fresh parsley, finely chopped": "parsley",
        "finely chopped parsley": "parsley",
        "fresh parsley leaves": "parsley",
        "fresh parsley": "parsley",
        "fresh flat leaf parsley leaves": "parsley",
        "fresh parsley, chopped": "parsley",
        "parsley,": "parsley",
        "leaves of fresh basil": "basil",
        "fresh basil": "basil",
        "leaves fresh basil": "basil",
        "basil pesto-1/3 fincan fesle\u011fen pesto": "basil",
        "sun dried tomatoes, chopped": "tomato",
        "tomatoes-grated": "tomato",
        "tomatoes": "tomato",
        "tomato puree": "tomato",
        "tomatoes, chopped": "tomato",
        "can crushed tomatoes": "tomato",
        "tomatoes, crushed": "tomato",
        "tomatoes,": "tomato",
        "tomatoes- chopped": "tomato",
        "vanilla extract\n-50 g": "vanilla",
        "vanilla": "vanilla",
        "vanilla bean-1 vanilya fasulyesi": "vanilla",
        "vanilla extract-5 ml vanilya \u00f6z\u00fc": "vanilla",
        "vanilla pod": "vanilla",
        "vanilla essence": "vanilla",
        "vanilla extract": "vanilla",
        "vanilla bean (or 1 tsp vanilla extract)": "vanilla",
        "heavy cream-250g a\u011f\u0131r krema": "cream",
        "single cream": "cream",
        "heavy cream 50 g": "cream",
        "cream, room temperature": "cream",
        "cream cheese 100 g": "cream",
        "cream cheese, room temperature": "cream",
        "coriander leaves": "coriander",
        "fresh coriander, chopped": "coriander",
        "fresh coriander leaves": "coriander",
        "chilli pepper": "chilli",
        "small chilli,de-seeded & finely chopped": "chilli",
        "chilli powder": "chilli",
        "chilli flakes": "chilli",
        "fresh chopped thyme": "thyme",
        "spring of thyme": "thyme",
        "of fresh thyme": "thyme",
        "fresh thyme and fresh tarragon": "thyme",
        "chopped fresh herbs (thyme, marjoram, sage)": "thyme",
        "thyme(optional)": "thyme",
        "dried thyme": "thyme",
        "fresh thyme and bay leaves": "thyme",
        "fresh thyme": "thyme",
        "of toasted almonds": "almond",
        "almonds,": "almond",
        "almonds-chopped-100 gr badem do\u011franm\u0131\u015f": "almond",
        "ground almonds": "almond",
        "oatmeal-": "oatmeal",
        "rolled oatmeal": "oatmeal",
        "oatmeal": "oatmeal",
        "active dried yeast  or 15g/1/2 oz fresh yeast": "yeast",
        "sachet dried yeast or 15g fresh yeast-7 gr po\u015fet kuru maya veya 15 gr taze maya": "yeast",
        "active dry yeast": "yeast",
        "active dried yeast  or 15g  fresh yeast": "yeast",
        "dried yeast or 15g fresh yeast": "yeast",
        "fresh yeast": "yeast",
        "balsamic vinegar": "vinegar",
        "vinegar": "vinegar",
        "finely ground -lean minced beef": "beef",
        "minced beef": "beef",
        "cubed beef fillet": "beef",
        "leg of pork (boneless)": "pork",
        "minced  pork": "pork",
        "pork chops": "pork",
        "baking soda-1 \u00e7ay ka\u015f\u0131\u011f\u0131 kabartma tozu": "baking powder",
        "baking powder": "baking powder",
        "whole turkey": "turkey",
        "turkey thigh": "turkey",
        "zest of an orange-1 orta boy portakal": "orange",
        "orange": "orange",
        "oranges \u2013cut into two and sliced": "orange",
        "zest of an orange\u00a0\u00a0-1 orta boy portakal kabu\u011fu rendesi": "orange",
        "zest of an orange": "orange",
        "baking potatoes": "potato",
        "potatoes, very thinly sliced": "potato",
        "potatoes": "potato",
        "lrg.potato,peeled&diced into 1cm cubes": "potato",
        "potatoes-": "potato",
        "limes, zest and juice-2 adet orta boy limon suyu": "lime",
        "lime, cut into thin slices": "lime",
        "juice of lime": "lime",
        "chicken drumsticks": "chicken",
        "boneless chicken breast , cut into bite-sized pieces": "chicken",
        "chicken portions": "chicken",
        "whole chicken": "chicken",
        "chicken breasts": "chicken",
        "mint leaves": "mint",
        "mint- finely chopped": "mint",
        "springs of rosemary": "rosemary",
        "of rosemary": "rosemary",
        "chopped spinach leaves": "spinach",
        "baby spinach- finely chopped": "spinach",
        "fresh spinach leaves, chopped": "spinach",
        "chopped spinach": "spinach",
        "ground nutmeg": "nutmeg",
        "nutmeg": "nutmeg",
        "fresh nutmeg, optional-": "nutmeg",
        "butter, melted": "butter",
        "unsalted butter": "butter",
        "butter, melted-90 gr tereya\u011f\u0131, eritilmi\u015f": "butter",
        "butter, cut into cubes": "butter",
        "butter, very cold": "butter",
        "butter, softened": "butter",
        "butter": "butter",
        "unsalted butter, chopped": "butter",
        "cold butter, diced": "butter",
        "ham, cut into stripes":"ham",
        "cocoa powder, for dusting":"cocoa",
        "sunflower seeds-50 gr ay\u00e7ekirde\u011fi":"sunflower",
        "sunflower seeds-": "sunflower",
        "shredded coconut": "coconut",
        "pinenut-roasted": "pinenut",
        "grated zucchini": "zucchini",
        "zucchini, sliced- dilimlenmi\u015f 1 kabak": "zucchini",
        "ground cinnamon-- 5 gr \u00f6\u011f\u00fct\u00fclm\u00fc\u015f tar\u00e7\u0131n": "cinnamon",
        "cinnamon": "cinnamon",
        "ground cinnamon": "cinnamon",
        "of cinnamon-": "cinnamon",
        "lamb": "lamb",
        "minced lean lamb": "lamb",
        "rack of lamb": "lamb",
        "bone-in leg of lamb": "lamb",
        "baby carrot": "carrot",
        "carrots, peeled and sliced": "carrot",
        "baby carrots": "carrot",
        "mussels": "mussels",
        "turkey": "turkey",
        "milk": "milk",
        "shallot, minced": "shallot",
        "shallot": "shallot",
        "tablespoons brandy, optional": "brandy",
        "rice": "rice",
        "asparagus, snap off and discard the woody ends.- 250 gr ku\u015fkonmaz, odunsu u\u00e7lar\u0131 kopar\u0131n ve at\u0131n.": "asparagus",
        "asparagus -500 gr ku\u015fkonmaz": "asparagus",
        "asparagus": "asparagus",
        "asparagus, cut into small pieces": "asparagus",
        "marinade": "marinade",
        "water": "water",
        "boiling water, to mix": "water",
        "warm water": "water",
        "warm water (38-43 c)": "water",
        "of chilled water": "water",
        "starch +20 ml water (to dissolve starch)": "water",
        "pint water": "water",
        "water, luke warm": "water",
        "chilled water": "water",
        "warm water- 300ml \u0131l\u0131k su": "water",
        "sesame seeds": "sesame seed",
        "sesame seed": "sesame seed",
        "honey": "honey",
        "honey or agave syrup": "honey",
        "of honey": "honey",
        "honey or agave": "honey",
        "honey-1/2 \u00e7ay ka\u015f\u0131\u011f\u0131 bal": "honey",
        "eggplants": "eggplant",
        "pistachio": "",
        "parmesan, grated": "parmesan",
        "parmesan cheese, grated": "parmesan",
        "crumbled feta cheese, divided-100 gr ufalanm\u0131\u015f beyaz peynir, b\u00f6l\u00fcnm\u00fc\u015f": "feta",
        "crumbled feta cheese, divided-100 gr ufalanm\u0131\u015f beyaz peynir, b\u00f6l\u00fcnm\u00fc\u015f": "feta",
        "feta": "feta",
        "feta cheese, shredded": "feta",
        "rolled oats": "oat",
        "rolled oats -300 gr yulaf ezmesi": "oat",
        "dried paprika": "paprika",
        "paprika": "paprika",
        "dried  paprika": "paprika",
        "smoked paprika": "paprika",
        "whole artichokes": "artichokes",
        "artichokes": "artichokes",
        "marinated artichokes drain and roughly chopped-350 gr marine edilmi\u015f enginar s\u00fcz\u00fcl\u00fcr ve kabaca do\u011fran\u0131r": "artichokes",
        "artichokes, chopped": "artichokes",
        "agave": "agave",
        "bay leaf": "bay",
        "bay leaves": "bay",
        "large shrimp": "shrimp",
        "of coconut oil-60 gr": "coconut oil",
        "coconut oil": "coconut oil",
        "coconut oil-60 gr hindistancevizi ya\u011f\u0131": "coconut oil",
        "whole duck": "duck",
        "yellow squashes": "squash",
        "dried figs, diced": "fig",
        "maple syrup-100 ml ak\u00e7aa\u011fa\u00e7 \u015furubu": "maple syrup",
        "maple syrup 25 g": "maple syrup",
        "toasted pine nuts-2 yemek ka\u015f\u0131\u011f\u0131 k\u0131zarm\u0131\u015f \u00e7am f\u0131st\u0131\u011f\u0131": "pine",
        "pinenut": "pine",
        "parmesan, grated": "parmesan",
        "parmesan cheese, grated": "parmesan",
        "raisins": "raisin",
        "of raisins": "raisin",
        "raw pumpkin seeds-50 gr \u00e7i\u011f kabak \u00e7ekirde\u011fi": "pumpkin seed",
        "pumpkin seeds": "pumpkin seed",
        "cheddar cheese, grated-": "cheddar",
        "cheddar cheese, grated": "cheddar",
        "capers, finely chopped": "capers",
        "shaved coconut-50 gr hindistan cevizi rendesi": "coconut",
        "coconut": "coconut",
        "warm-80 g \u0131l\u0131k s\u00fct": "",
        "corn starch": "",
        "good-quality apples": "apple",
        "apples, sliced": "apple",
        "apple puree-160 gr elma p\u00fcresi": "apple",
        "apple, thinly sliced   350 g": "apple",
        "apples-": "apple",
        "apple  100g": "apple",
        "apple juice to wet the oats": "apple",
        "grated mozzarella": "mozzarella",
        "shredded mozzarella": "mozzarella",
        "dried oregano": "oregano",
        "of fresh oregano": "oregano",
        "dried cumin": "cumin",
        "whole peppercorns": "pepper",
        "yellow bell pepper, diced": "bell pepper",
        "yellow bell pepper, diced": "bell pepper",
        "red bell pepper, diced": "bell pepper",
        "small pepper, deseeded and diced": "bell pepper",
        "freshly \u2013 cracked black pepper": "pepper",
        "walnuts, coarsely chopped": "walnut",
        "chopped walnuts": "walnut",
        "walnuts, finely chopped": "walnut",
        "finely chopped walnuts": "walnut",
        "poppy seeds": "poppy seed",
        "fresh dill": "dill",
        "of fresh dill": "dill",
        "dill, finely chopped": "dill",
        "bourbon": "bourbon",
        "all spice": "all spice",
        "olive oil-1 yemek ka\u015f\u0131\u011f\u0131 zeytinya\u011f\u0131": "olive oil",
        "olive oil-": "olive oil",
        "olive oil": "olive oil",
        "olive oil-2 yemek ka\u015f\u0131\u011f\u0131 zeytinya\u011f\u0131": "olive oil",
        "olive oil -1 yemek ka\u015f\u0131\u011f\u0131 zeytinya\u011f\u0131": "olive oil",
        "fish fillets (sea bass)": "bass",
        "whole sea bass": "bass",
        "fruits of your choice": "fruits",
        "graham cracker crumbs": "",
        "salmon fillet": "salmon",
        "boneless, skinless salmon": "salmon",
        "salmon fillets": "salmon",
        "burrata cheese (2-3 balls)- 250 gr burrata peyniri (2-3 top)": "burrata",
        "fresh strawberries, halved": "strawberry",
        "strawberries 500 g": "strawberry",
        "blueberries": "blueberriy",
        "dried blueberries-- 60 gr kurutulmu\u015f yaban mersini": "blueberry",
        "sirloin  or  rump steak": "steak",
        "sirloin steak": "steak",
        "fresh green salad-\u00bd fincan taze ye\u015fil salata": "green salad",
        "arugula -2 bardak roka": "arugula",
        "streaky bacon or panchetta slices": "bacon",
        "dried fennel seeds, crushed": "fennel",
        "fennel \u2013thinly sliced": "fennel",
        "sole fish": "sole",
        "white wine": "wine",
        "mascarpone cheese 100 g": "mascarpone",
        "eggs + 2 egg yolks, room temperature": "egg",
        "dried cayenne powder": "cayenne",
        "courgettes, peeled and sliced": "courgette",
        "of frozen puff pastry": "puff pastry",
        "dried apricots": "apricot",
        "ricotta cheese-250 gr ricotta peyniri": "ricotta",
        "fresh firm ricotta": "ricotta",
        "natural yoghurt": "yoghurt",
        "dried lasagne sheet": "lasagne",
        "hazelnuts": "hazelnut",
        "whole sea bream": "bream",
        "dates, mashed": "date",
        "of whole wheat bread, finely coarsed": "bread",
        "fresh whole wheat breadcrumbs": "bread",
        "fresh or frozen bluberry": "blueberry",
        "mulberries": "mulberry",
        "i\u0307lave tereya\u011f": "tereya",
        "pear   100g": "pear",
        "courgette- \u00a0thinly sliced lengthways\u00a0about 24 slices": "courgette",
        "bittersweet or semisweet chocolate, roughly chopped": "chocolate",
        "dark chocolate": "chocolate",
        "chia seeds": "chia seed",
        "banana, mashed": "banana",
        "chives, finely chopped": "chives",
        "phyllo dough 12 yaprak yufka hamuru": "phyllo dough",
        "frozen peas-thawed": "pea",
        "": "",
    }

    recipe["quantities"] = []
    recipe["ingredients"] = []
    unmatched_ingredient_labels = []

    for ingredient in recipe["original_ingredients"]:
        ingredient_parts_raw = ingredient.split(";")
        measurement_raw = ingredient_parts_raw[0].lower().strip()
        ingredient_label_raw = ingredient_parts_raw[1].lower().strip()

        if ingredient_label_raw in ingredient_label_agg_map:
            mapped_items = ingredient_label_agg_map[ingredient_label_raw]
            if isinstance(mapped_items, list):
                for mapped_item in mapped_items:
                    unmatched_ingredient_labels.append(mapped_item)
                    ingredient_label_raw = mapped_item
            else:
                ingredient_label_raw = mapped_items

        # handle ingredient labels
        if ingredient_label_raw not in ingr_labels_raw_to_r1m:
            unmatched_ingredient_labels.append(ingredient_label_raw)
            continue

        # TODO append the parsed ingredient list to the reicpe object


        # handle quantities
        measurement = regex.match(regex_measurements_with_fractions, ingredient)
        test1 = measurement.groups() if measurement else None
        # test = measurement[0]
        measurement2 = regex.match(regex_measurements_without_fractions, ingredient)
        # test2 = measurement[0]
        test2 = measurement2.groups() if measurement2 else None
        m1 = regex.search(regex_measurements_with_fractions, ingredient)
        test3 = m1.groups() if m1 else None
        m2 = regex.search(regex_measurements_without_fractions, ingredient)
        test4 = m2.groups() if m2 else None
        # TODO append the quantity to the recipe object

    return recipe, unmatched_ingredient_labels

def mapUnmatchedToR1m(unmatched_ingredient_labels, fg_labels):
    ingr_labels_raw_to_r1m = {}
    ingr_labels_r1m_to_raw = {}

    for unmatched_label in unmatched_ingredient_labels:
        if not unmatched_label:
            continue
        ingr_labels_raw_to_r1m[unmatched_label] = []
        for fg_label in fg_labels:
            if fg_label not in ingr_labels_r1m_to_raw:
                ingr_labels_r1m_to_raw[fg_label] = []
            if unmatched_label in fg_label:
                ingr_labels_raw_to_r1m[unmatched_label].append(fg_label)
                ingr_labels_r1m_to_raw[fg_label].append(unmatched_label)

    return ingr_labels_raw_to_r1m, ingr_labels_r1m_to_raw

def main():
    ARCELIC_RECIPES_DIR_PATH = os.path.abspath("./inputs/arcelik_recipe/Revize Edilen Tarifler/")
    GRAPH_NODES_PATH = os.path.abspath("./inputs/graph/nodes_191120.csv")

    RAW_TO_R1M_PATH = os.path.abspath("./outputs/arcelic_raw_to_r1m_unreviewed.json")
    R1M_TO_RAW_PATH = os.path.abspath("./outputs/arcelic_r1m_to_raw_unreviewed.json")


    unmatched_ingredient_labels = []

    fg_ingredients = get_graph_nodes(GRAPH_NODES_PATH)

    # TODO load from json
    ingr_labels_raw_to_r1m = {"salt": ["salt", "rock_salt"]}
    ingr_labels_r1m_to_raw = {"salt": ["salt"],
                              "rock_salt": ["salt"]}

    for file_name in os.listdir(ARCELIC_RECIPES_DIR_PATH):
        file_path = os.path.join(ARCELIC_RECIPES_DIR_PATH, file_name)
        recipe = parse_doc_to_recipe(file_path)
        recipe, _unmatched_ingredient_labels = parse_ingredients(recipe, ingr_labels_raw_to_r1m, ingr_labels_r1m_to_raw)
        unmatched_ingredient_labels += _unmatched_ingredient_labels



    print(len(unmatched_ingredient_labels))
    unmatched_ingredient_labels = list(set(unmatched_ingredient_labels))
    print(len(unmatched_ingredient_labels))
    # try normalizing with spacy
    # unmatched_ingredient_labels = [extract_main_keyword_singular(ingredient) for ingredient in unmatched_ingredient_labels]
    # print(len(unmatched_ingredient_labels))

    executeHardcodedIngredientMappings(unmatched_ingredient_labels)
    with open(os.path.abspath("./outputs/arcelic_raw_ingredient_labels.json"), "w") as json_file:
        json.dump(unmatched_ingredient_labels, json_file, indent=2)

    ingr_labels_raw_to_r1m, ingr_labels_r1m_to_raw = mapUnmatchedToR1m(unmatched_ingredient_labels, fg_ingredients)

    with open(RAW_TO_R1M_PATH, "w") as raw_to_r1m_file:
        json.dump(ingr_labels_raw_to_r1m, raw_to_r1m_file, indent=2)
    with open(R1M_TO_RAW_PATH, "w") as r1m_to_raw_file:
        json.dump(ingr_labels_r1m_to_raw, r1m_to_raw_file, indent=2)





if __name__ == "__main__":
    main()